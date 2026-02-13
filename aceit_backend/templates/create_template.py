"""
Script to create a professional resume template
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume_template():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('{{NAME}}')
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(31, 78, 121)
    
    # Contact Info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run('{{EMAIL}} | {{PHONE}}')
    contact_run.font.size = Pt(10)
    
    # LinkedIn and GitHub
    links_para = doc.add_paragraph()
    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_run = links_para.add_run('{{LINKEDIN}} | {{GITHUB}}')
    links_run.font.size = Pt(10)
    links_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Add spacing
    doc.add_paragraph()
    
    # Professional Summary Section
    summary_heading = doc.add_paragraph()
    summary_heading_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
    summary_heading_run.font.size = Pt(14)
    summary_heading_run.font.bold = True
    summary_heading_run.font.color.rgb = RGBColor(31, 78, 121)
    
    summary_content = doc.add_paragraph('{{SUMMARY}}')
    summary_content.paragraph_format.space_after = Pt(12)
    
    # Skills Section
    skills_heading = doc.add_paragraph()
    skills_heading_run = skills_heading.add_run('TECHNICAL SKILLS')
    skills_heading_run.font.size = Pt(14)
    skills_heading_run.font.bold = True
    skills_heading_run.font.color.rgb = RGBColor(31, 78, 121)
    
    skills_content = doc.add_paragraph('{{SKILLS}}')
    skills_content.paragraph_format.space_after = Pt(12)
    
    # Projects Section
    projects_heading = doc.add_paragraph()
    projects_heading_run = projects_heading.add_run('PROJECTS')
    projects_heading_run.font.size = Pt(14)
    projects_heading_run.font.bold = True
    projects_heading_run.font.color.rgb = RGBColor(31, 78, 121)
    
    doc.add_paragraph('{{PROJECTS}}')
    
    # Experience Section
    experience_heading = doc.add_paragraph()
    experience_heading_run = experience_heading.add_run('EXPERIENCE')
    experience_heading_run.font.size = Pt(14)
    experience_heading_run.font.bold = True
    experience_heading_run.font.color.rgb = RGBColor(31, 78, 121)
    
    doc.add_paragraph('{{EXPERIENCE}}')
    
    # Education Section
    education_heading = doc.add_paragraph()
    education_heading_run = education_heading.add_run('EDUCATION')
    education_heading_run.font.size = Pt(14)
    education_heading_run.font.bold = True
    education_heading_run.font.color.rgb = RGBColor(31, 78, 121)
    
    doc.add_paragraph('{{EDUCATION}}')
    
    # Certifications Section (Optional)
    certifications_heading = doc.add_paragraph()
    certifications_heading_run = certifications_heading.add_run('CERTIFICATIONS')
    certifications_heading_run.font.size = Pt(14)
    certifications_heading_run.font.bold = True
    certifications_heading_run.font.color.rgb = RGBColor(31, 78, 121)
    
    doc.add_paragraph('{{CERTIFICATIONS}}')
    
    # Save template
    doc.save('resume_template.docx')
    print("✅ Resume template created successfully!")

if __name__ == "__main__":
    create_resume_template()
